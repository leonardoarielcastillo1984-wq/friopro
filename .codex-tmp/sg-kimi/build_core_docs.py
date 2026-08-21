from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT=Path('outputs/sg-kimi-work/documentos_nuevos'); OUT.mkdir(parents=True,exist_ok=True)
BLUE='1F4E78'; LIGHT='D9EAF7'; GRAY='F2F4F7'

docs=[
('MC-01','Manual del Sistema de Gestión de la Calidad','A',['Contexto, partes interesadas y alcance','Mapa e interacción de procesos','Liderazgo, política y objetivos','Planificación de riesgos y oportunidades','Gestión de recursos y conocimiento','Operación y control de cambios','Evaluación del desempeño y mejora']),
('PG-01','Control de información documentada','A',['Creación, identificación y codificación','Revisión, aprobación y distribución','Control de cambios y versiones','Acceso, protección, respaldo y recuperación','Retención, disposición y documentos externos']),
('PG-02','Gestión de riesgos, oportunidades y contingencias','A',['Identificación y evaluación de riesgos','Planes de tratamiento','Planes de contingencia y pruebas','Continuidad de suministros, TI, personas e infraestructura','Revisión de eficacia y lecciones aprendidas']),
('PG-03','Auditorías internas del sistema y de procesos','A',['Programa basado en riesgo','Competencia e independencia de auditores','Auditorías de sistema, proceso y producto','Hallazgos, correcciones y acciones correctivas','Seguimiento y cierre de eficacia']),
('PG-04','No conformidades, solución de problemas y acciones correctivas','A',['Contención y protección del cliente','Análisis de causa raíz','Acción sistémica y prevención de recurrencia','Verificación de eficacia','Actualización de PFMEA, plan de control y lecciones aprendidas']),
('PG-05','Gestión de proveedores y procesos contratados','A',['Selección y evaluación inicial','Seguimiento de desempeño y escalamiento','Desarrollo del sistema de gestión del proveedor','Requisitos legales, reglamentarios y del cliente','Control de subproveedores y cambios']),
('PG-06','Competencia, formación y toma de conciencia','A',['Perfiles y competencias requeridas','Inducción y capacitación en el puesto','Evaluación de eficacia','Competencia de auditores e inspectores','Conciencia sobre calidad, seguridad del producto y consecuencias']),
('PG-07','Diseño y desarrollo de procesos y productos','A',['Planificación multidisciplinaria','Entradas y características especiales','Revisiones, verificación y validación','Salidas, cambios y trazabilidad','Desarrollo de procesos de fabricación y logística']),
('PG-08','Planificación y control operacional','A',['Revisión de requisitos del cliente','APQP, factibilidad y planificación','PFMEA, plan de control e instrucciones estandarizadas','Liberación, trazabilidad, preservación y entrega','Control de cambios temporales y permanentes']),
('PG-09','Mantenimiento, infraestructura y herramental','A',['Mantenimiento preventivo y predictivo','Equipos, herramental y repuestos críticos','Objetivos de disponibilidad y OEE cuando aplique','Conservación, almacenamiento y propiedad del cliente','Planes de reacción ante fallas']),
('PG-10','Seguimiento, medición, calibración y laboratorio','A',['Programa de calibración y verificación','Trazabilidad metrológica','Laboratorios internos y externos','Estudios MSA','Control de equipos no conformes']),
('PG-11','Seguridad del producto y requisitos específicos del cliente','A',['Responsabilidades y escalamiento','Características especiales y controles','Aprobaciones especiales y trazabilidad','Gestión de cambios y comunicación al cliente','Lecciones aprendidas y requisitos específicos']),
('PG-12','Revisión por la dirección y objetivos de calidad','A',['Entradas obligatorias y desempeño de procesos','Satisfacción del cliente y scorecards','Costos de mala calidad, riesgos y recursos','Decisiones, acciones y responsables','Seguimiento de objetivos y eficacia del sistema']),
]

def shade(cell,color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),color); tcPr.append(shd)

def build(code,title,rev,sections):
    d=Document(); sec=d.sections[0]; sec.top_margin=sec.bottom_margin=Inches(0.8); sec.left_margin=sec.right_margin=Inches(0.85)
    styles=d.styles; normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
    for s,size,col in [('Title',24,BLUE),('Heading 1',16,BLUE),('Heading 2',12,'1F4D78')]:
        st=styles[s]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(col); st.font.bold=True
    h=sec.header.paragraphs[0]; h.text='SISTEMA DE GESTIÓN KIMI  |  ISO 9001:2015 + IATF 16949'; h.style=styles['Normal']; h.runs[0].font.size=Pt(8); h.runs[0].font.color.rgb=RGBColor.from_string('68717A')
    f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; f.add_run(f'{code} | Revisión {rev} | Copia no controlada si se imprime').font.size=Pt(8)
    p=d.add_paragraph(); p.style='Title'; p.add_run(title)
    t=d.add_table(rows=2,cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    vals=[('Código',code),('Revisión',rev),('Estado','Pendiente de aprobación'),('Fecha','19/08/2026')]
    for i,(a,b) in enumerate(vals):
        c1=t.cell(i//2,(i%2)*2); c2=t.cell(i//2,(i%2)*2+1); c1.text=a; c2.text=b; shade(c1,LIGHT); c1.paragraphs[0].runs[0].bold=True
        c1.vertical_alignment=c2.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    d.add_heading('1. Objeto y alcance',1)
    d.add_paragraph(f'Establecer los criterios obligatorios de KIMI para {title.lower()}, aplicables a los procesos, sedes, productos, servicios y proveedores comprendidos en el alcance del sistema de gestión. Este documento integra los requisitos pertinentes de ISO 9001:2015 e IATF 16949 y los requisitos específicos del cliente que resulten contractualmente aplicables.')
    d.add_heading('2. Referencias y definiciones',1)
    for x in ['ISO 9001:2015 — Sistemas de gestión de la calidad.','IATF 16949 — Requisitos para organizaciones de producción automotriz y piezas de servicio.','Requisitos legales, reglamentarios, contractuales y específicos del cliente vigentes.','Maestro de documentos, matriz de riesgos, mapa de procesos y registros relacionados.']:
        d.add_paragraph(x,style='List Bullet')
    d.add_heading('3. Responsabilidades',1)
    for x in ['Dirección: aprobar, proveer recursos y revisar la eficacia.','Calidad: mantener este documento, verificar cumplimiento y conservar evidencias.','Responsables de proceso: implementar controles, medir resultados y tratar desvíos.','Todo el personal: cumplir instrucciones, informar riesgos y preservar registros.']:
        d.add_paragraph(x,style='List Bullet')
    d.add_heading('4. Requisitos del sistema',1)
    for i,s in enumerate(sections,1):
        d.add_heading(f'4.{i} {s}',2)
        d.add_paragraph(f'El responsable del proceso deberá definir, aplicar y mantener controles documentados para {s.lower()}. Los criterios de aceptación, responsabilidades, recursos, riesgos, métodos de seguimiento y registros deberán estar identificados antes de ejecutar la actividad. Todo desvío se contendrá, analizará y cerrará con evidencia de eficacia.')
    d.add_heading('5. Gestión de riesgos y cambios',1)
    d.add_paragraph('Antes de introducir cambios se evaluará su impacto sobre conformidad, seguridad del producto, capacidad, continuidad, trazabilidad, requisitos del cliente y documentos relacionados. Los cambios deberán ser aprobados por las funciones autorizadas y, cuando corresponda, por el cliente.')
    d.add_heading('6. Indicadores y seguimiento',1)
    for x in ['Cumplimiento del plan y de los criterios definidos.','Cantidad, antigüedad y recurrencia de desvíos.','Eficacia de acciones y cumplimiento de plazos.','Impacto en cliente, producto, entrega, costo y seguridad.']:
        d.add_paragraph(x,style='List Bullet')
    d.add_heading('7. Registros y conservación',1)
    d.add_paragraph('Los registros deberán ser legibles, identificables, protegidos contra pérdida o alteración y recuperables durante el plazo definido en el maestro documental, requisitos del cliente y legislación aplicable. Las evidencias electrónicas deberán incluir respaldo y control de acceso.')
    d.add_heading('8. Control de cambios',1)
    tab=d.add_table(rows=2,cols=4); tab.style='Table Grid'; tab.rows[0].cells[0].text='Revisión'; tab.rows[0].cells[1].text='Fecha'; tab.rows[0].cells[2].text='Descripción'; tab.rows[0].cells[3].text='Aprobación'
    for c in tab.rows[0].cells: shade(c,BLUE); c.paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255); c.paragraphs[0].runs[0].bold=True
    tab.rows[1].cells[0].text=rev; tab.rows[1].cells[1].text='19/08/2026'; tab.rows[1].cells[2].text='Emisión para integración ISO 9001:2015 / IATF 16949.'; tab.rows[1].cells[3].text='Pendiente'
    d.add_paragraph('\nElaboró: ____________________    Revisó: ____________________    Aprobó: ____________________')
    d.save(OUT/f'{code}_{title.replace(" ","_").replace("/","-")}_Rev_{rev}.docx')

for x in docs: build(*x)
print(f'generados={len(docs)}')
