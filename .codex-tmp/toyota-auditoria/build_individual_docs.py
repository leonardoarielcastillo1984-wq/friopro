from pathlib import Path
import re
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

import build_compliance_doc as m

ROOT = m.BASE / 'outputs/toyota-documentos-individuales'
CATEGORY_FOLDERS = {
    'Normativa': '01_Normativa',
    'Método y Gestión': '02_Metodo_y_Gestion',
    'Factor humano': '03_Factor_Humano',
    'Entorno': '04_Entorno',
    'Tecnología': '05_Tecnologia',
    'Máquina': '06_Maquina',
}


def safe_name(value):
    value = re.sub(r'[^A-Za-z0-9ÁÉÍÓÚÜÑáéíóúüñ _-]+', '', value)
    value = re.sub(r'\s+', '_', value.strip())
    return value[:85]


def setup_doc(doc, item, code):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.68)
    sec.left_margin = Inches(0.78)
    sec.right_margin = Inches(0.78)
    sec.header_distance = Inches(0.32)
    sec.footer_distance = Inches(0.32)

    normal = doc.styles['Normal']
    normal.font.name = 'Aptos'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos')
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(m.BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in [
        ('Title', 23, m.BLUE, 0, 7),
        ('Heading 1', 16, m.BLUE, 12, 6),
        ('Heading 2', 12.5, m.MID_BLUE, 9, 4),
        ('Heading 3', 11.2, m.BLUE, 6, 3),
    ]:
        st = doc.styles[name]
        st.font.name = 'Aptos Display'
        st._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos Display')
        st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos Display')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True
    for list_name in ['List Bullet', 'List Bullet 2']:
        st = doc.styles[list_name]
        st.font.name = 'Aptos'
        st.font.size = Pt(10.2)
        st.paragraph_format.space_after = Pt(3)
        st.paragraph_format.line_spacing = 1.12

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(f'DADA 2026  |  {code}')
    m.font_run(r, size=8.5, bold=True, color=m.GRAY)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run('Documento controlado · Página ')
    m.font_run(r, size=8, color=m.GRAY)
    m.add_field(fp, 'PAGE')


def add_title_block(doc, item, dtype, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item['category'].upper())
    m.font_run(r, size=10, bold=True, color=m.GOLD)
    p = doc.add_paragraph(style='Title')
    r = p.add_run(item['point'])
    m.font_run(r, size=23, bold=True, color=m.BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(dtype)
    m.font_run(r, size=12.5, color=m.MID_BLUE)
    m.add_metadata_table(doc, [
        ('Código', code),
        ('Versión', '1.0 · Borrador para aprobación'),
        ('Responsable', item['dept']),
        ('Participan', m.contributors(item['dept'], item['category'])),
        ('Vigencia', 'Desde su aprobación hasta la próxima revisión'),
        ('Revisión', 'Anual y ante cambios o incidentes relevantes'),
    ])


def add_scope(doc, item):
    doc.add_heading('1. Objeto y alcance', level=1)
    m.add_label_paragraph(doc, 'Pregunta de auditoría:', item['question'])
    subject = item['point'].split('–', 1)[-1].strip().lower()
    p = doc.add_paragraph()
    r = p.add_run(
        f'Este documento establece los criterios obligatorios para gestionar {subject} en todas las operaciones, bases, '
        'unidades, conductores, personal propio y terceros alcanzados por la actividad de transporte de DADA.'
    )
    m.font_run(r)


def add_policy(doc, item):
    doc.add_heading('2. Lineamientos obligatorios', level=1)
    for clause in m.policy_clauses(item['point'], item['question'], item['dept'], item['category']):
        m.add_bullet(doc, clause)
    m.add_bullet(doc, 'Todo incumplimiento deberá documentarse, evaluarse según su riesgo y tratarse mediante acciones correctivas con plazo y responsable.')
    m.add_bullet(doc, 'Los registros generados deberán ser legibles, recuperables, protegidos y conservarse durante el plazo legal o contractual aplicable.')


def add_roles(doc, item):
    doc.add_heading('3. Responsabilidades', level=1)
    role_table = doc.add_table(rows=1, cols=2)
    m.set_table_geometry(role_table, [2500, 6860])
    for j, text in enumerate(['Rol', 'Responsabilidad']):
        m.set_cell_shading(role_table.rows[0].cells[j], m.MID_BLUE)
        r = role_table.rows[0].cells[j].paragraphs[0].add_run(text)
        m.font_run(r, size=9, bold=True, color=m.WHITE)
    m.set_repeat_table_header(role_table.rows[0])
    roles = [
        ('Dirección', 'Aprobar el documento, asignar recursos y revisar el desempeño.'),
        (item['dept'], 'Implementar el proceso, conservar evidencias, informar desvíos y medir resultados.'),
        ('Calidad', 'Controlar versiones, auditar la aplicación y verificar la eficacia de las acciones.'),
        ('Personal alcanzado', 'Cumplir el estándar, participar de capacitaciones y reportar anormalidades.'),
    ]
    for idx, (role, responsibility) in enumerate(roles):
        row = role_table.add_row()
        if idx % 2:
            m.set_cell_shading(row.cells[0], m.LIGHT_GRAY)
            m.set_cell_shading(row.cells[1], m.LIGHT_GRAY)
        for j, value in enumerate([role, responsibility]):
            r = row.cells[j].paragraphs[0].add_run(value)
            m.font_run(r, size=9.2, bold=(j == 0), color=m.BLUE if j == 0 else m.BLACK)


def add_procedure(doc, item):
    doc.add_heading('4. Procedimiento de implementación', level=1)
    num_id = m.create_numbering_instance(doc)
    for step in m.implementation_steps(item['point'], item['question'], item['evidence'], item['dept'], item['category']):
        m.add_numbered_step(doc, step, num_id)

    doc.add_heading('5. Registros obligatorios', level=1)
    for record in m.records_from_evidence(item['evidence'], item['point']):
        m.add_bullet(doc, record)

    doc.add_heading('6. Indicadores y revisión', level=1)
    for metric in m.kpis(item['point'], item['category']):
        m.add_bullet(doc, metric)
    m.add_bullet(doc, 'Frecuencia mínima de revisión: mensual para la operación y anual para el documento.')


def add_forms(doc, item, code):
    doc.add_page_break()
    doc.add_heading('ANEXO A · Registro de verificación', level=1)
    m.add_metadata_table(doc, [
        ('Código relacionado', code),
        ('Fecha de verificación', '____ / ____ / ______'),
        ('Centro / operación', '____________________________________________'),
        ('Unidad / persona', '____________________________________________'),
        ('Verificador', '____________________________________________'),
    ])

    records = m.records_from_evidence(item['evidence'], item['point'])
    table = doc.add_table(rows=1, cols=5)
    m.set_table_geometry(table, [4600, 760, 760, 760, 2480])
    for j, text in enumerate(['Evidencia o control', 'Sí', 'No', 'N/A', 'Ubicación / referencia']):
        m.set_cell_shading(table.rows[0].cells[j], m.MID_BLUE)
        r = table.rows[0].cells[j].paragraphs[0].add_run(text)
        m.font_run(r, size=8.5, bold=True, color=m.WHITE)
    m.set_repeat_table_header(table.rows[0])
    for idx, record in enumerate(records):
        row = table.add_row()
        if idx % 2:
            for cell in row.cells:
                m.set_cell_shading(cell, m.LIGHT_GRAY)
        values = [record, '☐', '☐', '☐', '']
        for j, value in enumerate(values):
            p = row.cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j in (1, 2, 3) else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            m.font_run(r, size=8.8)

    doc.add_heading('Resultado de la verificación', level=2)
    result = doc.add_table(rows=4, cols=2)
    m.set_table_geometry(result, [2700, 6660])
    for i, (label, value) in enumerate([
        ('Resultado', '☐ Conforme     ☐ Parcial     ☐ No conforme     ☐ No aplica'),
        ('Porcentaje de cumplimiento', '________ %'),
        ('Observaciones', '\n\n'),
        ('Evidencia adjunta', '\n'),
    ]):
        m.set_cell_shading(result.rows[i].cells[0], m.LIGHT_BLUE)
        r = result.rows[i].cells[0].paragraphs[0].add_run(label)
        m.font_run(r, size=9, bold=True, color=m.BLUE)
        r = result.rows[i].cells[1].paragraphs[0].add_run(value)
        m.font_run(r, size=9)

    doc.add_heading('ANEXO B · Plan de acción', level=1)
    actions = doc.add_table(rows=1, cols=6)
    m.set_table_geometry(actions, [700, 2800, 1400, 1250, 1250, 1960])
    for j, text in enumerate(['N.º', 'Desvío / acción', 'Responsable', 'Fecha', 'Estado', 'Evidencia de cierre']):
        m.set_cell_shading(actions.rows[0].cells[j], m.MID_BLUE)
        r = actions.rows[0].cells[j].paragraphs[0].add_run(text)
        m.font_run(r, size=8, bold=True, color=m.WHITE)
    m.set_repeat_table_header(actions.rows[0])
    for idx in range(1, 6):
        row = actions.add_row()
        for j, value in enumerate([str(idx), '\n', '', '', 'Pendiente', '']):
            r = row.cells[j].paragraphs[0].add_run(value)
            m.font_run(r, size=8.5)

    doc.add_heading('Aprobación y cierre', level=2)
    approval = doc.add_table(rows=2, cols=3)
    m.set_table_geometry(approval, [3120, 3120, 3120])
    for j, label in enumerate(['Elaboró', 'Revisó', 'Aprobó']):
        m.set_cell_shading(approval.rows[0].cells[j], m.LIGHT_BLUE)
        r = approval.rows[0].cells[j].paragraphs[0].add_run(label)
        m.font_run(r, size=9, bold=True, color=m.BLUE)
        r = approval.rows[1].cells[j].paragraphs[0].add_run('Nombre / firma / fecha\n\n')
        m.font_run(r, size=8.5, italic=True, color=m.GRAY)


def build_one(item):
    dtype = m.document_type(item['point'], item['question'])
    code = m.code_for(int(item['n']), dtype)
    doc = Document()
    setup_doc(doc, item, code)
    add_title_block(doc, item, dtype, code)
    add_scope(doc, item)
    add_policy(doc, item)
    add_roles(doc, item)
    add_procedure(doc, item)
    add_forms(doc, item, code)
    folder = ROOT / CATEGORY_FOLDERS[item['category']]
    folder.mkdir(parents=True, exist_ok=True)
    path = folder / f"{int(item['n']):02d}_{safe_name(item['point'])}.docx"
    doc.save(path)
    return path


if ROOT.exists():
    shutil.rmtree(ROOT)
ROOT.mkdir(parents=True)
paths = [build_one(item) for item in m.items]

# Package-level reference files.
shutil.copy2(m.OUT, ROOT / '00_Carpeta_Maestra_Cumplimiento_DADA_2026.docx')
shutil.copy2(m.XLSX, ROOT / '00_Plan_Auditoria_Transporte_DADA_2026.xlsx')

index_doc = Document()
setup_doc(index_doc, m.items[0], 'IND-SV-00')
p = index_doc.add_paragraph(style='Title')
r = p.add_run('Índice del paquete documental Toyota')
m.font_run(r, size=23, bold=True, color=m.BLUE)
p = index_doc.add_paragraph()
r = p.add_run('83 documentos individuales organizados por categoría · DADA 2026')
m.font_run(r, size=12.5, color=m.MID_BLUE)
idx_table = index_doc.add_table(rows=1, cols=4)
m.set_table_geometry(idx_table, [700, 2500, 3560, 2600])
for j, text in enumerate(['N.º', 'Código / punto', 'Documento', 'Responsable']):
    m.set_cell_shading(idx_table.rows[0].cells[j], m.MID_BLUE)
    rr = idx_table.rows[0].cells[j].paragraphs[0].add_run(text)
    m.font_run(rr, size=8.5, bold=True, color=m.WHITE)
m.set_repeat_table_header(idx_table.rows[0])
for item, path in zip(m.items, paths):
    row = idx_table.add_row()
    for j, value in enumerate([str(item['n']), item['point'], path.name, item['dept']]):
        rr = row.cells[j].paragraphs[0].add_run(value)
        m.font_run(rr, size=8.2, bold=(j == 1), color=m.BLUE if j == 1 else m.BLACK)
index_doc.save(ROOT / '00_Indice_Paquete_Documental.docx')

print(f'GENERATED={len(paths)}')
print(ROOT)
