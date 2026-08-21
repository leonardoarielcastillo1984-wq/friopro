from pathlib import Path
from datetime import date
from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path('/Users/leonardocastillo/Desktop/APP/SGI respaldo 360')
XLSX = BASE / 'outputs/toyota-auditoria-2026/Plan_Auditoria_Transporte_DADA_2026.xlsx'
OUT = BASE / 'outputs/toyota-auditoria-2026/Carpeta_Maestra_Cumplimiento_DADA_2026.docx'

BLUE = '1F4E78'
MID_BLUE = '2F75B5'
LIGHT_BLUE = 'D9EAF7'
PALE_BLUE = 'EEF5FA'
GRAY = '5B6573'
LIGHT_GRAY = 'F2F4F7'
GOLD = 'C69214'
WHITE = 'FFFFFF'
BLACK = '202124'
GREEN = '548235'
RED = '9C0006'


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for tag, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(total))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    layout = tbl_pr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def font_run(run, size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = 'Aptos'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Aptos')
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Aptos')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field_code
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.extend([fld_char1, instr, fld_char2])


def add_label_paragraph(doc, label, text, color=BLACK, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(label + ' ')
    font_run(r, bold=True, color=BLUE)
    r = p.add_run(text)
    font_run(r, color=color)
    return p


def add_bullet(doc, text, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    font_run(r, size=10.2)
    return p


def create_numbering_instance(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))]
    num_ids = [int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(abstract_id))
    multi = OxmlElement('w:multiLevelType')
    multi.set(qn('w:val'), 'singleLevel')
    abstract.append(multi)
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    num_fmt = OxmlElement('w:numFmt')
    num_fmt.set(qn('w:val'), 'decimal')
    lvl_text = OxmlElement('w:lvlText')
    lvl_text.set(qn('w:val'), '%1.')
    suff = OxmlElement('w:suff')
    suff.set(qn('w:val'), 'tab')
    p_pr = OxmlElement('w:pPr')
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'num')
    tab.set(qn('w:pos'), '540')
    tabs.append(tab)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '540')
    ind.set(qn('w:hanging'), '280')
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '60')
    spacing.set(qn('w:line'), '280')
    spacing.set(qn('w:lineRule'), 'auto')
    p_pr.extend([tabs, ind, spacing])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_step(doc, text, num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    num_id_el = OxmlElement('w:numId')
    num_id_el.set(qn('w:val'), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    r = p.add_run(text)
    font_run(r, size=10.2)
    return p


def document_type(point, question):
    s = f'{point} {question}'.lower()
    if any(k in s for k in ['política', 'politica', 'compromiso', 'objetivos']):
        return 'Política y declaración de compromiso'
    if any(k in s for k in ['manual', 'normas de manejo', 'manejo defensivo']):
        return 'Manual / estándar operativo'
    if any(k in s for k in ['capacitación', 'capacitacion', 'charlas', 'inducción', 'induccion']):
        return 'Programa de capacitación y registro'
    if any(k in s for k in ['mantenimiento', 'calibración', 'calibracion', 'inspección', 'inspeccion', 'vida útil', 'vida util']):
        return 'Procedimiento técnico y registro de flota'
    if any(k in s for k in ['emergencia', 'accidente', 'incidente', 'antipánico', 'antipanico']):
        return 'Procedimiento / protocolo de respuesta'
    if any(k in s for k in ['evaluación', 'evaluacion', 'auditoría', 'auditoria', 'verificación', 'verificacion']):
        return 'Procedimiento de evaluación y formulario de registro'
    if any(k in s for k in ['certificación', 'certificacion', 'legal', 'normativo']):
        return 'Matriz de cumplimiento y expediente de respaldo'
    if any(k in s for k in ['sistema', 'monitoreo', 'gps', 'cámara', 'camara', 'sensor', 'rastreo']):
        return 'Estándar de operación tecnológica y registro'
    return 'Procedimiento documentado y registro de evidencia'


def code_for(n, doc_type):
    prefix = 'POL' if doc_type.startswith('Política') else 'MAN' if doc_type.startswith('Manual') else 'CAP' if doc_type.startswith('Programa') else 'PR'
    return f'{prefix}-SV-{n:02d}'


def contributors(dept, category):
    base = {
        'Seguridad e Higiene': 'Tráfico, RR. HH., Seguimiento, Mantenimiento y Calidad, según corresponda.',
        'RR. HH.': 'Seguridad e Higiene, Tráfico y responsables jerárquicos.',
        'Tráfico': 'Seguridad e Higiene, Seguimiento, Mantenimiento y conductores.',
        'Seguimiento': 'Tráfico, Seguridad e Higiene, Tecnología y Mantenimiento.',
        'Mantenimiento e Ingeniería de Flota': 'Tráfico, Seguridad e Higiene, Seguimiento y Contrataciones.',
        'Calidad/Sistema de Gestión': 'Dirección y responsables de los procesos auditados.',
        'Dirección/Gerencia General': 'Todos los departamentos involucrados.',
        'Contrataciones': 'Legales, Mantenimiento, Seguridad e Higiene y Tráfico.',
        'Legales': 'Calidad, Seguridad e Higiene y Dirección.',
        'Seguridad Patrimonial': 'Seguimiento, Tráfico, Legales y Seguridad e Higiene.',
    }
    return base.get(dept, 'Áreas operativas y de soporte relacionadas con el requisito.')


def policy_clauses(point, question, dept, category):
    subject = point.split('–', 1)[-1].strip()
    s = f'{point} {question}'.lower()
    clauses = [
        f'DADA establece como obligatorio gestionar {subject.lower()} mediante criterios documentados, responsables definidos y evidencias trazables.',
        f'El departamento {dept} será propietario del proceso, conservará los registros y reportará los desvíos a la Gerencia y a Calidad.',
    ]
    if category == 'Tecnología':
        clauses.append('La solución deberá contar con inventario de cobertura, prueba funcional, mantenimiento, gestión de alertas y protocolo ante indisponibilidad.')
    elif category == 'Máquina':
        clauses.append('Ninguna unidad podrá operar cuando exista una falla crítica de seguridad; toda liberación deberá quedar autorizada y registrada.')
    elif category == 'Factor humano':
        clauses.append('La competencia y aptitud del personal se validarán antes de la habilitación y periódicamente, preservando la confidencialidad de los datos personales.')
    elif category == 'Normativa':
        clauses.append('Los requisitos legales y corporativos serán revisados al menos anualmente y ante todo cambio normativo, contractual u operativo relevante.')
    else:
        clauses.append('El proceso se revisará mediante indicadores, auditorías y acciones correctivas bajo un ciclo de mejora continua.')
    if any(k in s for k in ['subcontrat', 'proveedor']):
        clauses.append('Los terceros estarán sujetos, como mínimo, a los mismos estándares exigidos a la operación propia y a controles previos y periódicos.')
    if any(k in s for k in ['fatiga', 'somnolencia', 'descanso', 'jornada']):
        clauses.append('Ante señales de fatiga, exceso de jornada o descanso insuficiente, se interrumpirá la asignación o conducción hasta recuperar una condición segura.')
    if any(k in s for k in ['accidente', 'incidente', 'emergencia']):
        clauses.append('Todo evento será comunicado inmediatamente, contenido, investigado y cerrado con análisis de causa, contramedidas y difusión de aprendizajes.')
    return clauses


def implementation_steps(point, question, evidence, dept, category):
    subject = point.split('–', 1)[-1].strip()
    steps = [
        f'Designar formalmente al responsable de {subject.lower()} y aprobar este documento mediante firma de la autoridad correspondiente.',
        'Comunicar el contenido a las personas alcanzadas y registrar asistencia, entrega o aceptación digital.',
        f'Implementar el control operativo bajo responsabilidad de {dept}, utilizando formatos con fecha, responsable, resultado, desvío y acción.',
        'Revisar mensualmente los registros y elevar desvíos críticos de forma inmediata; realizar una revisión integral al menos una vez por año.',
    ]
    if category in ('Tecnología', 'Máquina'):
        steps.insert(2, 'Verificar en una muestra representativa de unidades el funcionamiento real, la cobertura, el mantenimiento y el plan alternativo ante falla.')
    if category == 'Factor humano':
        steps.insert(2, 'Validar competencia y comprensión mediante evaluación teórica, práctica u observación en campo, según el riesgo.')
    return steps


def records_from_evidence(evidence, point):
    chunks = [x.strip(' .;\n') for x in evidence.replace('\n', '. ').split('.') if x.strip()]
    records = []
    for chunk in chunks[:5]:
        if len(chunk) > 12:
            records.append(chunk[0].upper() + chunk[1:] + '.')
    if not records:
        records = [f'Registro de implementación y verificación de {point.split("–", 1)[-1].strip().lower()}.']
    records.extend(['Registro de desvíos, acciones correctivas y verificación de eficacia.', 'Control de versión, aprobación y evidencia de comunicación del documento.'])
    return records[:6]


def kpis(point, category):
    subject = point.split('–', 1)[-1].strip().lower()
    if category == 'Tecnología':
        return [f'Cobertura operativa de {subject}: unidades habilitadas / unidades alcanzadas × 100.', 'Alertas tratadas dentro del plazo / alertas totales × 100.']
    if category == 'Máquina':
        return [f'Cumplimiento del control de {subject}: controles realizados / controles programados × 100.', 'Fallas críticas abiertas y tiempo promedio de cierre.']
    if category == 'Factor humano':
        return [f'Personal alcanzado por {subject}: personas conformes / personas alcanzadas × 100.', 'Desvíos de competencia o conducta cerrados dentro del plazo × 100.']
    return [f'Cumplimiento de {subject}: evidencias conformes / evidencias requeridas × 100.', 'Acciones vencidas / acciones totales y porcentaje de cierre eficaz.']


def add_metadata_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [2350, 7010])
    for i, (label, value) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        set_cell_shading(c1, LIGHT_BLUE)
        for p in c1.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(label)
            font_run(r, size=9.5, bold=True, color=BLUE)
        for p in c2.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            font_run(r, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


wb = load_workbook(XLSX, data_only=False)
ws = wb['Plan de auditoría']
items = []
for r in range(5, 88):
    items.append({
        'n': ws.cell(r, 1).value,
        'category': ws.cell(r, 2).value,
        'point': ws.cell(r, 3).value,
        'question': ws.cell(r, 4).value or '',
        'evidence': ws.cell(r, 5).value or '',
        'dept': ws.cell(r, 6).value,
    })
assert len(items) == 83

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.78)
section.right_margin = Inches(0.78)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Aptos'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Aptos')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Aptos')
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, before, after in [
    ('Title', 27, BLUE, 0, 8),
    ('Heading 1', 17, BLUE, 14, 7),
    ('Heading 2', 13.5, MID_BLUE, 10, 5),
    ('Heading 3', 11.5, BLUE, 7, 3),
]:
    st = styles[name]
    st.font.name = 'Aptos Display' if name != 'Normal' else 'Aptos'
    st._element.rPr.rFonts.set(qn('w:ascii'), st.font.name)
    st._element.rPr.rFonts.set(qn('w:hAnsi'), st.font.name)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ['List Bullet', 'List Bullet 2', 'List Number']:
    st = styles[list_name]
    st.font.name = 'Aptos'
    st.font.size = Pt(10.2)
    st.paragraph_format.space_after = Pt(3)
    st.paragraph_format.line_spacing = 1.12

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hp.add_run('DADA 2026  |  Sistema de Seguridad del Transporte')
font_run(r, size=8.5, bold=True, color=GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('Documento modelo sujeto a aprobación e implementación  ·  Página ')
font_run(r, size=8, color=GRAY)
add_field(fp, 'PAGE')

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(92)
p.paragraph_format.space_after = Pt(8)
r = p.add_run('CARPETA MAESTRA')
font_run(r, size=13, bold=True, color=GOLD)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph(style='Title')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sistema documental de cumplimiento')
font_run(r, size=27, bold=True, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(28)
r = p.add_run('Auditoría de Seguridad del Transporte · DADA 2026')
font_run(r, size=15, color=MID_BLUE)

cover_table = doc.add_table(rows=4, cols=2)
set_table_geometry(cover_table, [2600, 6760])
for i, (label, value) in enumerate([
    ('Organización', 'DADA'),
    ('Alcance', '83 requisitos de la autoevaluación Toyota'),
    ('Versión', 'Borrador 1.0 para revisión y aprobación'),
    ('Fecha', date.today().strftime('%d/%m/%Y')),
]):
    c1, c2 = cover_table.rows[i].cells
    set_cell_shading(c1, LIGHT_BLUE)
    r = c1.paragraphs[0].add_run(label)
    font_run(r, bold=True, color=BLUE)
    r = c2.paragraphs[0].add_run(value)
    font_run(r)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(28)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Uso interno · Documento controlado')
font_run(r, size=9.5, italic=True, color=GRAY)
doc.add_page_break()

# Front matter
doc.add_heading('Cómo utilizar esta carpeta', level=1)
p = doc.add_paragraph()
r = p.add_run('Propósito. ')
font_run(r, bold=True, color=BLUE)
r = p.add_run('Proporcionar un modelo documental completo para implementar y demostrar cada punto de la auditoría. Cada ficha incluye el documento requerido, cláusulas modelo, pasos de implementación, registros e indicadores.')
font_run(r)

note = doc.add_table(rows=1, cols=1)
set_table_geometry(note, [9360])
set_cell_shading(note.cell(0, 0), 'FFF2CC')
p = note.cell(0, 0).paragraphs[0]
r = p.add_run('IMPORTANTE · ')
font_run(r, bold=True, color='7F6000')
r = p.add_run('La existencia del documento no demuestra por sí sola el cumplimiento. Debe completarse con datos reales, aprobarse, comunicarse, implementarse y respaldarse con registros, certificados, pruebas de campo o tecnología, según corresponda.')
font_run(r, color='7F6000')

doc.add_heading('Reglas de control documental', level=2)
for text in [
    'Completar los campos de aprobación, vigencia, revisión y responsables antes de emitir cada documento.',
    'Conservar las versiones obsoletas identificadas y evitar su utilización operativa.',
    'Vincular cada registro con el punto de auditoría, fecha, unidad, persona o proceso correspondiente.',
    'Proteger datos médicos, personales y disciplinarios mediante permisos de acceso y anonimización.',
    'Revisar el sistema al menos anualmente y ante incidentes, cambios legales, tecnológicos u operativos.',
]:
    add_bullet(doc, text)

doc.add_heading('Mapa documental recomendado', level=2)
summary_table = doc.add_table(rows=1, cols=4)
set_table_geometry(summary_table, [900, 2600, 3260, 2600])
headers = ['N.º', 'Punto', 'Documento propuesto', 'Responsable']
for j, text in enumerate(headers):
    set_cell_shading(summary_table.rows[0].cells[j], MID_BLUE)
    r = summary_table.rows[0].cells[j].paragraphs[0].add_run(text)
    font_run(r, size=8.5, bold=True, color=WHITE)
set_repeat_table_header(summary_table.rows[0])
for item in items:
    row = summary_table.add_row()
    dtype = document_type(item['point'], item['question'])
    vals = [str(item['n']), item['point'], dtype, item['dept']]
    for j, value in enumerate(vals):
        if item['n'] % 2 == 0:
            set_cell_shading(row.cells[j], LIGHT_GRAY)
        r = row.cells[j].paragraphs[0].add_run(value)
        font_run(r, size=8.2, bold=(j == 1), color=BLUE if j == 1 else BLACK)

doc.add_page_break()

current_category = None
for index, item in enumerate(items):
    n = int(item['n'])
    category = item['category']
    point = item['point']
    question = item['question']
    evidence = item['evidence']
    dept = item['dept']
    dtype = document_type(point, question)
    code = code_for(n, dtype)

    if category != current_category:
        if current_category is not None:
            doc.add_page_break()
        current_category = category
        doc.add_heading(category.upper(), level=1)
        p = doc.add_paragraph()
        r = p.add_run('Documentos modelo y controles correspondientes a esta categoría de la auditoría.')
        font_run(r, italic=True, color=GRAY)

    doc.add_heading(f'{n}. {point}', level=2)
    add_metadata_table(doc, [
        ('Documento requerido', dtype),
        ('Código propuesto', code),
        ('Responsable principal', dept),
        ('Participan', contributors(dept, category)),
        ('Aprobación requerida', 'Gerencia del área y Dirección / Gerencia General'),
        ('Revisión', 'Anual y ante cambios, incidentes o desvíos relevantes'),
    ])

    add_label_paragraph(doc, 'Pregunta de auditoría:', question)
    add_label_paragraph(doc, 'Objetivo del documento:', f'Establecer criterios verificables para responder y demostrar el cumplimiento de {point.split("–", 1)[-1].strip().lower()}, asegurando aplicación sistemática, registros y mejora continua.')

    doc.add_heading('Contenido normativo modelo', level=3)
    for clause in policy_clauses(point, question, dept, category):
        add_bullet(doc, clause)

    doc.add_heading('Implementación obligatoria', level=3)
    step_num_id = create_numbering_instance(doc)
    for step in implementation_steps(point, question, evidence, dept, category):
        add_numbered_step(doc, step, step_num_id)

    doc.add_heading('Registros y evidencia a conservar', level=3)
    for record in records_from_evidence(evidence, point):
        add_bullet(doc, record)

    doc.add_heading('Indicadores mínimos', level=3)
    for metric in kpis(point, category):
        add_bullet(doc, metric)

    doc.add_heading('Criterio de cierre del punto', level=3)
    closure = (
        'El requisito podrá declararse implementado cuando el documento esté aprobado y vigente; el personal alcanzado haya sido comunicado o capacitado; '
        'existan registros reales y trazables; los desvíos cuenten con acciones y responsables; y una verificación independiente confirme la eficacia en campo.'
    )
    p = doc.add_paragraph()
    r = p.add_run(closure)
    font_run(r)

    approval = doc.add_table(rows=2, cols=3)
    set_table_geometry(approval, [3120, 3120, 3120])
    for j, label in enumerate(['Elaboró', 'Revisó', 'Aprobó']):
        set_cell_shading(approval.rows[0].cells[j], LIGHT_BLUE)
        r = approval.rows[0].cells[j].paragraphs[0].add_run(label)
        font_run(r, size=9, bold=True, color=BLUE)
        r = approval.rows[1].cells[j].paragraphs[0].add_run('Nombre / firma / fecha')
        font_run(r, size=8.5, italic=True, color=GRAY)

    if index != len(items) - 1:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Closing annex
doc.add_page_break()
doc.add_heading('Anexo · Lista de verificación previa a auditoría', level=1)
for text in [
    'Cada uno de los 83 puntos tiene un propietario y una persona asignada.',
    'Todos los documentos tienen código, versión, aprobación, fecha de vigencia y próxima revisión.',
    'Las evidencias corresponden a operaciones reales y no solamente a modelos vacíos.',
    'Los registros médicos y personales están anonimizados o protegidos.',
    'Los sistemas tecnológicos fueron probados en campo y cuentan con mantenimiento y contingencia.',
    'Los desvíos abiertos tienen responsable, plazo y seguimiento de eficacia.',
    'Los subcontratados están incluidos en el mismo estándar de gestión.',
    'La Dirección conoce los resultados, recursos necesarios y principales brechas.',
]:
    add_bullet(doc, text)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
