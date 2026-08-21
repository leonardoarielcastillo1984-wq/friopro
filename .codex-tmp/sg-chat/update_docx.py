from pathlib import Path
from docx import Document
from datetime import datetime, timezone
import json,re

ROOT=Path('/Users/leonardocastillo/Desktop/SG chat')
items={x['ruta']:x for x in json.loads(Path('outputs/sg-chat-work/contenido_extraido.json').read_text(encoding='utf-8'))}
modules={
'control_documental':('ISO 9001:2015 7.5 / IATF 16949 7.5.3.2.1','La información documentada debe identificarse, revisarse, aprobarse, protegerse, mantenerse disponible en el punto de uso y retirarse cuando quede obsoleta. Los cambios deberán conservar trazabilidad de revisión, fecha, responsable y motivo.'),
'recursos_humanos':('ISO 9001:2015 7.1.2, 7.2 y 7.3 / IATF 16949 7.2.1–7.3.2','La organización deberá definir competencias, proporcionar formación, evaluar su eficacia y conservar evidencia. El personal deberá conocer su contribución a la calidad, seguridad del producto y consecuencias del incumplimiento.'),
'proveedores':('ISO 9001:2015 8.4 / IATF 16949 8.4.1–8.4.3','Los proveedores y procesos contratados se seleccionarán, controlarán y reevaluarán según riesgo y desempeño. Deberán comunicarse requisitos técnicos, legales, del cliente, de cambios, trazabilidad y seguridad del producto.'),
'operacion':('ISO 9001:2015 8.1 y 8.5 / IATF 16949 8.5.1.1–8.5.1.7','La ejecución deberá realizarse bajo condiciones controladas, con criterios de aceptación, instrucciones vigentes, recursos adecuados, trazabilidad y registros. Los cambios temporales o permanentes requieren evaluación, autorización y verificación posterior.'),
'calidad':('ISO 9001:2015 8.6, 8.7, 9.1, 9.2 y 10.2 / IATF 16949 9.2.2 y 10.2.3','La liberación y el seguimiento deberán conservar evidencia de conformidad y autorización. Los desvíos requieren identificación, contención, causa raíz, acción sistémica y verificación documentada de eficacia.'),
'riesgos':('ISO 9001:2015 6.1 / IATF 16949 6.1.2.1 y 6.1.2.3','Los riesgos y oportunidades deberán evaluarse según probabilidad, impacto y controles existentes. Las contingencias críticas deberán contar con respuesta, responsables, comunicación, pruebas periódicas y revisión de resultados.'),
'cliente':('ISO 9001:2015 8.2 y 9.1.2 / IATF 16949 9.1.2.1','Se identificarán y revisarán requisitos contractuales, legales y específicos del cliente antes de asumir compromisos. La satisfacción, reclamos y scorecards deberán analizarse y generar acciones cuando no se alcancen los objetivos.'),
'metrologia':('ISO 9001:2015 7.1.5 / IATF 16949 7.1.5.1','Los recursos de medición deberán ser aptos, identificados y trazables a patrones reconocidos. Se controlarán calibraciones, verificaciones, MSA, laboratorios y el impacto de equipos fuera de condición.'),
'mantenimiento':('ISO 9001:2015 7.1.3 / IATF 16949 8.5.1.5 y 8.5.1.6','La infraestructura, equipos y herramentales críticos deberán mantenerse mediante planes preventivos y predictivos, repuestos, objetivos de disponibilidad y planes de reacción frente a fallas.'),
'logistica':('ISO 9001:2015 8.5.4 y 8.6 / IATF 16949 8.5.4.1','La manipulación, almacenamiento, embalaje, transporte y entrega deberán preservar la conformidad, identificación y trazabilidad. Los desvíos de entrega o daño requieren reacción, comunicación y registro.'),
}

def next_rev(s):
    return chr(ord(s.upper())+1) if len(s)==1 and s.isalpha() and s.upper()<'Z' else s

count=0; failures=[]
for p in sorted(ROOT.rglob('*.docx')):
    rel=str(p.relative_to(ROOT)); item=items.get(rel,{})
    try:
        d=Document(p)
        # Update visible single-letter revision markers without flattening runs.
        for para in list(d.paragraphs)+[p for t in d.tables for row in t.rows for c in row.cells for p in c.paragraphs]:
            for run in para.runs:
                run.text=re.sub(r'(?i)(revisi[oó]n\s*[:.-]?\s*)([A-Z])\b',lambda m:m.group(1)+next_rev(m.group(2)),run.text)
        d.add_paragraph()
        h=d.add_paragraph('Actualización del sistema de gestión — 19/08/2026')
        try: h.style='Heading 1'
        except Exception: h.runs[0].bold=True
        d.add_paragraph('Esta actualización complementa y precisa los controles del documento sin modificar su finalidad operativa ni su diseño original. Su aplicación queda sujeta a aprobación del responsable del proceso y de la autoridad definida en el maestro documental.')
        kinds=item.get('clasificacion') or ['control_documental']
        for k in kinds:
            if k not in modules: continue
            ref,txt=modules[k]
            q=d.add_paragraph(); r=q.add_run(ref+': '); r.bold=True; q.add_run(txt)
        d.add_paragraph('Registros mínimos: evidencia de ejecución, responsable, fecha, resultado, desvíos, acciones y verificación de cierre, según corresponda al propósito de este documento.')
        d.core_properties.modified=datetime.now(timezone.utc)
        try: d.core_properties.revision=int(d.core_properties.revision or 0)+1
        except Exception: d.core_properties.revision=1
        d.save(p); count+=1
    except Exception as e: failures.append((rel,str(e)))
Path('outputs/sg-chat-work/docx_failures.json').write_text(json.dumps(failures,ensure_ascii=False,indent=2),encoding='utf-8')
print('updated',count,'failed',len(failures))
